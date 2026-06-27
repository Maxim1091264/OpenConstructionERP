"""Adapter for DOCX-based document understanding."""

from __future__ import annotations

from typing import Optional
from uuid import uuid4

try:
    import docx
except ImportError:  # pragma: no cover
    docx = None  # type: ignore[assignment]

from app.modules.ai_smeta_ru.domain.models import ExtractedObject


class DocxAdapter:
    """Extracts paragraphs and tables from DOCX files."""

    def extract(self, file_path: str, artifact_id: str | None = None) -> list[ExtractedObject]:
        if docx is None:
            raise ImportError(
                "DOCX extraction requires python-docx. Install it with: pip install python-docx"
            )

        document = docx.Document(file_path)
        artifact_id = artifact_id or str(uuid4())
        objects: list[ExtractedObject] = []

        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            object_id = f"docx-paragraph-{index}-{uuid4().hex}"
            objects.append(
                ExtractedObject(
                    object_id=object_id,
                    artifact_id=artifact_id,
                    object_type="paragraph",
                    raw_text=text,
                    normalized_text=text,
                    confidence=0.85,
                )
            )

        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if not values:
                    continue
                raw_text = " | ".join(values)
                object_id = f"docx-table-{table_index}-{row_index}-{uuid4().hex}"
                objects.append(
                    ExtractedObject(
                        object_id=object_id,
                        artifact_id=artifact_id,
                        object_type="table_row",
                        raw_text=raw_text,
                        normalized_text=raw_text,
                        confidence=0.9,
                    )
                )

        return objects
